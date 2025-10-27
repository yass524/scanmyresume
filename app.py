import os
import json
import uuid
import time
import secrets
import hmac, hashlib, base64
from io import BytesIO
from pathlib import Path
from typing import Optional
from contextlib import asynccontextmanager

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse
from pydantic import BaseModel

from sqlalchemy import create_engine, Column, Integer, String, DateTime, func
from sqlalchemy.orm import sessionmaker, declarative_base
from passlib.context import CryptContext

# Optional: embeddings (guarded by EMB_ON)
try:
    import torch  # noqa: F401
except Exception:
    torch = None

# ---------- Optional deps (PDF, parsing) ----------
try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    _PDF_OK = True
except Exception:
    _PDF_OK = False

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.layout import LAParams
except Exception:
    pdf_extract_text, LAParams = None, None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from pdf2image import convert_from_bytes as pdf2img_from_bytes
except Exception:
    pdf2img_from_bytes = None

# DOCX (used for .docx resumes / JDs)
try:
    import docx  # python-docx
except Exception:
    docx = None

# ---------- Local modules ----------
from ats_core import compute_score, _maybe_load_model

# ---------- Paths & config ----------
HERE         = Path(__file__).resolve().parent
INDEX_PATH   = HERE / "index.html"
LOGIN_PATH   = HERE / "login.html"
REGISTER_PATH= HERE / "register.html"
HOME_PATH    = HERE / "home.html"

# Detect serverless/container (Cloud Run sets K_SERVICE and PORT)
IS_SERVERLESS = bool(os.environ.get("K_SERVICE") or os.environ.get("PORT"))

# Reports directory: default to /tmp on serverless, local folder otherwise (can override with REPORT_DIR env)
REPORT_DIR = Path(os.environ.get("REPORT_DIR", "/tmp/reports" if IS_SERVERLESS else str(HERE / "reports")))
REPORT_DIR.mkdir(parents=True, exist_ok=True)

APP_TITLE = "ATS-like Resume Checker API"
VERSION   = "0.3.4"

# Upload limits
MAX_UPLOAD_MB = float(os.environ.get("MAX_UPLOAD_MB", "8"))
MAX_BYTES     = int(MAX_UPLOAD_MB * 1024 * 1024)

# Rate limit (req/min per IP)
MAX_RPM = int(os.environ.get("MAX_RPM", "120"))
_REQ_LOG: dict[str, list[float]] = {}

# Auth config
LOGIN_ENABLED  = os.environ.get("LOGIN_ENABLED", "1") == "1"
DEMO_USER_RAW  = os.environ.get("DEMO_USER")
DEMO_PASS_RAW  = os.environ.get("DEMO_PASS")
DEMO_USER      = DEMO_USER_RAW.strip().lower() if isinstance(DEMO_USER_RAW, str) else None
DEMO_PASS      = DEMO_PASS_RAW.strip() if isinstance(DEMO_PASS_RAW, str) else None
DEMO_DEMO_ENABLED = bool(DEMO_USER and DEMO_PASS)
SESSION_COOKIE = "ats_session"

# Cookies: default secure on serverless/https, off locally (can override with env)
_COOKIE_SEC_DEFAULT = "1" if IS_SERVERLESS else "0"
COOKIE_SECURE  = bool(int(os.environ.get("COOKIE_SECURE", _COOKIE_SEC_DEFAULT)))

# ======= SESSION SECRET (updated: no warning) =======
# Production requires SESSION_SECRET; local dev will auto-generate a per-run secret.
_session_secret_env = os.environ.get("SESSION_SECRET")
if not _session_secret_env:
    if IS_SERVERLESS:
        raise RuntimeError("SESSION_SECRET env var required when running in serverless/production mode.")
    _session_secret_env = secrets.token_urlsafe(32)
SESSION_SECRET = _session_secret_env
# ====================================================

# DB (SQLite by default). On serverless, default to /tmp for writeable FS.
# Ensure Windows paths are valid SQLite URLs by using forward slashes.
if IS_SERVERLESS:
    _default_sqlite = "sqlite:////tmp/users.db"
else:
    _db_path = (HERE / "users.db").resolve()
    _default_sqlite = f"sqlite:///{_db_path.as_posix()}"
DB_URL = os.environ.get("DB_URL", _default_sqlite)
engine = create_engine(
    DB_URL,
    connect_args={"check_same_thread": False} if DB_URL.startswith("sqlite") else {}
)
SessionLocal = sessionmaker(bind=engine, autocommit=False, autoflush=False)
Base = declarative_base()

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# ---------- DB Models ----------
class User(Base):
    __tablename__ = "users"
    id            = Column(Integer, primary_key=True, index=True)
    username      = Column(String(50), unique=True, index=True, nullable=False)
    password_hash = Column(String(255), nullable=False)
    created_at    = Column(DateTime(timezone=True), server_default=func.now())

def hash_password(p: str) -> str:
    return pwd_context.hash(p)

def verify_password(p: str, h: str) -> bool:
    try:
        return pwd_context.verify(p, h)
    except Exception:
        return False

# ---------- Pydantic ----------
class ScoreRequest(BaseModel):
    resume_text: str
    job_description: str

# ---------- Lifespan ----------
@asynccontextmanager
async def lifespan(app: FastAPI):
    try:
        Base.metadata.create_all(engine)
    except Exception as e:
        print("DB init failed:", e)
    try:
        if os.environ.get("EMB_ON", "0") == "1":
            _maybe_load_model()
            print("Embeddings model preloaded.")
    except Exception as e:
        print("Embeddings preload skipped/failed:", e)
    yield

# ---------- Create app ----------
app = FastAPI(title=APP_TITLE, version=VERSION, lifespan=lifespan)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],      # tighten for prod
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- Helpers ----------
def rate_limit(request: Request):
    now = time.time()
    ip  = request.client.host if request.client else "unknown"
    q = _REQ_LOG.get(ip, [])
    q = [t for t in q if now - t < 60]  # drop old entries
    if len(q) >= MAX_RPM:
        raise HTTPException(429, "Too many requests, slow down.")
    q.append(now)
    _REQ_LOG[ip] = q

def _make_token(username: str) -> str:
    payload = (username or "user").encode("utf-8")
    sig = hmac.new(SESSION_SECRET.encode("utf-8"), payload, hashlib.sha256).hexdigest()
    data = base64.urlsafe_b64encode(payload).decode("utf-8").rstrip("=")
    return f"{data}.{sig}"

def _verify_token(token: str) -> bool:
    try:
        data, sig = token.split(".", 1)
        payload = base64.urlsafe_b64decode(data + "===")  # pad
        expected = hmac.new(SESSION_SECRET.encode("utf-8"), payload, hashlib.sha256).hexdigest()
        return hmac.compare_digest(sig, expected)
    except Exception:
        return False

def is_authed(request: Request) -> bool:
    if not LOGIN_ENABLED:
        return True
    tok = request.cookies.get(SESSION_COOKIE)
    return bool(tok) and _verify_token(tok)

def read_docx_bytes(b: bytes) -> str:
    if docx is None:
        raise HTTPException(500, "DOCX parsing requires 'python-docx'. Install it.")
    d = docx.Document(BytesIO(b))
    return "\n".join(p.text for p in d.paragraphs)

def try_pymupdf_extract_text(b: bytes) -> str:
    if fitz is None:
        return ""
    try:
        doc = fitz.open(stream=b, filetype="pdf")
        texts = []
        for page in doc:
            txt = page.get_text("text")
            if txt:
                texts.append(txt)
        return "\n".join(texts).strip()
    except Exception:
        return ""

def try_pdfminer_extract_text(b: bytes) -> str:
    if pdf_extract_text is None:
        return ""
    try:
        laparams = LAParams()
        return pdf_extract_text(BytesIO(b), laparams=laparams) or ""
    except Exception:
        return ""

def try_ocr_extract_text(b: bytes) -> str:
    if pytesseract is None or pdf2img_from_bytes is None:
        return ""
    try:
        pages = pdf2img_from_bytes(b, dpi=200)
        txts = []
        for img in pages:
            txts.append(pytesseract.image_to_string(img))
        return "\n".join(txts).strip()
    except Exception:
        return ""

def read_pdf_bytes(b: bytes) -> str:
    # Strongest first: PyMuPDF -> pdfminer -> OCR
    text = try_pymupdf_extract_text(b)
    if text and text.strip():
        return text
    text = try_pdfminer_extract_text(b)
    if text and text.strip():
        return text
    text = try_ocr_extract_text(b)
    return text

def sanitize_unicode(s: str) -> str:
    if not s:
        return s
    repl = {
        "’": "'", "‘": "'", "“": '"', "”": '"', "…": "...", "≥": ">=", "≤": "<=", "×": "x",
        "\u00a0": " ", "\u200b": "", "\uf0b7": "-", "\uf0a7": "-", "\uf02d": "-",
    }
    for k, v in repl.items():
        s = s.replace(k, v)
    return s

def ensure_size(b: bytes):
    if len(b) > MAX_BYTES:
        raise HTTPException(413, f"File too large (>{MAX_UPLOAD_MB} MB).")

def docx_stats_from_bytes(b: bytes) -> Optional[dict]:
    if docx is None:
        return None
    try:
        d = docx.Document(BytesIO(b))
        paras = [p for p in d.paragraphs]
        runs  = sum(len(p.runs) for p in paras)
        words = sum(len((p.text or "").split()) for p in paras)
        return {"paragraphs": len(paras), "runs": runs, "approx_words": words}
    except Exception:
        return None

def ext_from_filename(name: str) -> str:
    p = name.rsplit(".", 1)
    if len(p) == 2:
        return p[1].lower()
    return ""

# ---------- Static UI ----------
@app.get("/")
def home():
    if HOME_PATH.exists():
        return FileResponse(HOME_PATH)
    return JSONResponse({"ok": True, "service": APP_TITLE, "version": VERSION})

@app.get("/ui")
def index(request: Request):
    if LOGIN_ENABLED and not is_authed(request):
        # Redirect to login with message and next param
        return RedirectResponse(url="/login?next=%2Fui&msg=Please%20log%20in%20to%20continue", status_code=302)
    if INDEX_PATH.exists():
        return FileResponse(INDEX_PATH)
    return JSONResponse({"ok": True, "msg": "UI not bundled"})

@app.get("/login", include_in_schema=False)
def login_page():
    if LOGIN_PATH.exists():
        return FileResponse(LOGIN_PATH)
    return RedirectResponse(url="/ui", status_code=302)

@app.get("/register", include_in_schema=False)
def register_page():
    if REGISTER_PATH.exists():
        return FileResponse(REGISTER_PATH)
    return RedirectResponse(url="/ui", status_code=302)

# ---------- Auth ----------
@app.post("/auth/register")
def auth_register(username: str = Form(...), password: str = Form(...)):
    if not LOGIN_ENABLED:
        raise HTTPException(400, "Registration disabled")
    u = (username or "").strip().lower()
    p = (password or "").strip()
    if not u or not p:
        raise HTTPException(400, "Username and password required")
    db = SessionLocal()
    try:
        existing = db.query(User).filter_by(username=u).first()
        if existing:
            raise HTTPException(409, "Username already exists")
        user = User(username=u, password_hash=hash_password(p))
        db.add(user)
        db.commit()
        return {"ok": True}
    finally:
        db.close()

@app.post("/auth/login")
def auth_login(response: Response, username: str = Form(...), password: str = Form(...)):
    if not LOGIN_ENABLED:
        return {"ok": True, "redirect": "/ui"}
    u = username.strip().lower()
    db = SessionLocal()
    try:
        user = db.query(User).filter_by(username=u).first()
    finally:
        db.close()
    ok = False
    if user and verify_password(password, user.password_hash):
        ok = True
    if not ok:
        raise HTTPException(401, "Invalid credentials")
    tok = _make_token(u)
    resp = JSONResponse({"ok": True, "redirect": "/ui"})
    resp.set_cookie(SESSION_COOKIE, tok, httponly=True, samesite="lax", secure=COOKIE_SECURE, max_age=60*60*24*7)
    return resp

@app.post("/auth/logout")
def auth_logout():
    resp = JSONResponse({"ok": True})
    resp.delete_cookie(SESSION_COOKIE)
    return resp

@app.get("/health", include_in_schema=False)
def health():
    return {"ok": True, "version": VERSION}

# ---------- Reports ----------
@app.get("/r/{rid}")
def get_report(rid: str):
    p = REPORT_DIR / f"{rid}.json"
    if not p.exists():
        raise HTTPException(404, "Report not found")
    return FileResponse(p, media_type="application/json")

# ---------- Score helpers ----------
def build_warnings(report: dict, ext: str, r_bytes: bytes, docx_stats: Optional[dict]) -> list[str]:
    warns = []
    if ext == "pdf":
        warns.append("PDF detected. If text extraction seems incomplete, try saving as DOCX.")
    if docx_stats and docx_stats.get("approx_words", 0) > 1200:
        warns.append("Resume looks quite long; consider trimming to ~1 page (early career) or 2 pages (experienced).")
    if not report.get("action_verbs"):
        warns.append("Consider adding more strong action verbs (led, built, optimized, automated, …).")
    if report.get("contact_warnings"):
        warns.extend(report["contact_warnings"])
    return warns
# ---------- PDF helpers ----------
def _build_pdf_from_report(report: dict, out_path: Path):
    if not _PDF_OK:
        raise HTTPException(500, "PDF export not available. Install 'reportlab'.")
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    story = []

    # Title
    title = f"ATS Report • Score: {report.get('score', 0)}/100"
    story.append(Paragraph(title, styles['Title']))
    story.append(Spacer(1, 12))

    # Components table
    comps = report.get("components", {})
    if comps:
        data = [["Component", "Value"]] + [[k.replace("_", " "), str(v)] for k, v in comps.items()]
        t = Table(data, hAlign='LEFT')
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0f172a")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("GRID", (0,0), (-1,-1), 0.25, colors.gray),
            ("BACKGROUND", (0,1), (-1,-1), colors.HexColor("#111827")),
            ("TEXTCOLOR", (0,1), (-1,-1), colors.white),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#111827"), colors.HexColor("#0b1220")]),
            ("LEFTPADDING", (0,0), (-1,-1), 6),
            ("RIGHTPADDING", (0,0), (-1,-1), 6),
        ]))
        story.append(Paragraph("Components", styles['Heading2']))
        story.append(t)
        story.append(Spacer(1, 12))

    # Keywords
    def para_list(title, arr):
        story.append(Paragraph(title, styles['Heading2']))
        if arr:
            story.append(Paragraph(", ".join(arr), styles['BodyText']))
        else:
            story.append(Paragraph("None", styles['BodyText']))
        story.append(Spacer(1, 8))

    para_list("Matched Keywords", report.get("matched_keywords") or [])
    para_list("Missing Keywords", report.get("missing_keywords") or [])
    para_list("AI Matched Terms", report.get("ai_matched_terms") or [])

    # Suggestions
    sugg = report.get("suggested_bullets") or []
    story.append(Paragraph("Suggested Bullets", styles['Heading2']))
    if sugg:
        for s in sugg:
            story.append(Paragraph(f"• {s}", styles['BodyText']))
    else:
        story.append(Paragraph("None", styles['BodyText']))
    story.append(Spacer(1, 12))

    # Save
    doc = SimpleDocTemplate(str(out_path), pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    doc.build(story)

# ---------- PDF endpoint ----------
@app.get("/report/{rid}.pdf")
def get_report_pdf(rid: str):
    p_json = REPORT_DIR / f"{rid}.json"
    if not p_json.exists():
        raise HTTPException(404, "Report not found")
    p_pdf = REPORT_DIR / f"{rid}.pdf"

    # (Re)build PDF if missing or JSON is newer
    try:
        if (not p_pdf.exists()) or (p_pdf.stat().st_mtime < p_json.stat().st_mtime):
            report = json.loads(p_json.read_text("utf-8"))
            _build_pdf_from_report(report, p_pdf)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Failed to build PDF: {e}")

    return FileResponse(p_pdf, media_type="application/pdf", filename="ATS-report.pdf")

# ---------- Endpoints ----------
@app.post("/score-text")
def score_text(request: Request, payload: ScoreRequest):
    if not is_authed(request):
        raise HTTPException(401, "Login required")
    rate_limit(request)
    resume_text = sanitize_unicode((payload.resume_text or "").strip())
    jd_text     = sanitize_unicode((payload.job_description or "").strip())
    if not jd_text:
        raise HTTPException(400, "job_description is required.")
    if not resume_text:
        raise HTTPException(400, "resume_text is empty.")
    report = compute_score(resume_text, jd_text)
    report["format_warnings"] = build_warnings(report, ext="txt", r_bytes=b"", docx_stats=None)
    rid = str(uuid.uuid4())
    (REPORT_DIR / f"{rid}.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), "utf-8")
    report["id"] = rid
    report["share_url"] = f"/r/{rid}"
    return report

@app.post("/score", include_in_schema=False)
def score_legacy(request: Request, payload: ScoreRequest):
    """
    Backwards-compatible alias used by older front-end bundles.
    """
    return score_text(request, payload)

@app.post("/score-file")
async def score_file(
    request: Request,
    resume: UploadFile = File(...),
    job_description: UploadFile | None = File(None),
    jd_text: str | None = Form(None),
):
    if not is_authed(request):
        raise HTTPException(401, "Login required")
    rate_limit(request)

    r_bytes = await resume.read()
    if len(r_bytes) == 0:
        raise HTTPException(400, "Empty resume file.")
    ensure_size(r_bytes)
    r_name  = resume.filename or "resume"
    r_ext   = ext_from_filename(r_name)

    # JD source
    if job_description is not None:
        jd_bytes = await job_description.read()
        ensure_size(jd_bytes)
        jd_ext = ext_from_filename(job_description.filename or "jd")
        if jd_ext == "pdf":
            jd_text2 = read_pdf_bytes(jd_bytes)
        elif jd_ext == "docx":
            jd_text2 = read_docx_bytes(jd_bytes)
        else:
            try:
                jd_text2 = jd_bytes.decode("utf-8", errors="ignore")
            except Exception:
                jd_text2 = ""
    else:
        jd_text2 = jd_text or ""

    jd_text2 = sanitize_unicode((jd_text2 or "").strip())

    # Resume text
    resume_text = ""
    if r_ext == "pdf":
        resume_text = read_pdf_bytes(r_bytes)
    elif r_ext == "docx":
        resume_text = read_docx_bytes(r_bytes)
    else:
        try:
            resume_text = r_bytes.decode("utf-8", errors="ignore")
        except Exception:
            resume_text = ""

    resume_text = sanitize_unicode((resume_text or "").strip())

    if not jd_text2 or not jd_text2.strip():
        raise HTTPException(400, "Provide job_description file or jd_text.")
    if not resume_text.strip():
        raise HTTPException(400, "No text found in resume. If it’s a scanned PDF, convert to DOCX or use OCR.")

    report = compute_score(resume_text, jd_text2)
    docx_meta = docx_stats_from_bytes(r_bytes) if r_ext == "docx" else None
    report["format_warnings"] = build_warnings(report, ext=r_ext, r_bytes=r_bytes, docx_stats=docx_meta)

    rid = str(uuid.uuid4())
    (REPORT_DIR / f"{rid}.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), "utf-8")
    report["id"] = rid
    report["share_url"] = f"/r/{rid}"
    return report
from fastapi.responses import PlainTextResponse

@app.get("/ads.txt")
def ads_txt():
    return PlainTextResponse(
        "google.com, pub-8738953870753140, DIRECT, f08c47fec0942fa0"
    )
