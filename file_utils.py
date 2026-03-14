from io import BytesIO
from typing import Optional

from fastapi import HTTPException

from config import MAX_BYTES, MAX_UPLOAD_MB

try:
    import fitz
except Exception:
    fitz = None

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.layout import LAParams
except Exception:
    pdf_extract_text, LAParams = None, None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from pdf2image import convert_from_bytes as pdf2img_from_bytes
except Exception:
    pdf2img_from_bytes = None

try:
    import docx
except Exception:
    docx = None


def read_docx_bytes(b: bytes) -> str:
    if docx is None:
        raise HTTPException(500, "DOCX parsing requires 'python-docx'. Install it.")
    d = docx.Document(BytesIO(b))
    return "\n".join(p.text for p in d.paragraphs)


def try_pymupdf_extract_text(b: bytes) -> str:
    if fitz is None:
        return ""
    try:
        doc = fitz.open(stream=b, filetype="pdf")
        texts = []
        for page in doc:
            txt = page.get_text("text")
            if txt:
                texts.append(txt)
        return "\n".join(texts).strip()
    except Exception:
        return ""


def try_pdfminer_extract_text(b: bytes) -> str:
    if pdf_extract_text is None:
        return ""
    try:
        laparams = LAParams()
        return pdf_extract_text(BytesIO(b), laparams=laparams) or ""
    except Exception:
        return ""


def try_ocr_extract_text(b: bytes) -> str:
    if pytesseract is None or pdf2img_from_bytes is None:
        return ""
    try:
        pages = pdf2img_from_bytes(b, dpi=200)
        txts = []
        for img in pages:
            txts.append(pytesseract.image_to_string(img))
        return "\n".join(txts).strip()
    except Exception:
        return ""


def read_pdf_bytes(b: bytes) -> str:
    text = try_pymupdf_extract_text(b)
    if text and text.strip():
        return text
    text = try_pdfminer_extract_text(b)
    if text and text.strip():
        return text
    return try_ocr_extract_text(b)


def sanitize_unicode(s: str) -> str:
    if not s:
        return s
    repl = {
        "’": "'",
        "‘": "'",
        "“": '"',
        "”": '"',
        "…": "...",
        "≥": ">=",
        "≤": "<=",
        "×": "x",
        "\u00a0": " ",
        "\u200b": "",
        "\uf0b7": "-",
        "\uf0a7": "-",
        "\uf02d": "-",
    }
    for k, v in repl.items():
        s = s.replace(k, v)
    return s


def ensure_size(b: bytes):
    if len(b) > MAX_BYTES:
        raise HTTPException(413, f"File too large (>{MAX_UPLOAD_MB} MB).")


def docx_stats_from_bytes(b: bytes) -> Optional[dict]:
    if docx is None:
        return None
    try:
        d = docx.Document(BytesIO(b))
        paras = [p for p in d.paragraphs]
        runs = sum(len(p.runs) for p in paras)
        words = sum(len((p.text or "").split()) for p in paras)
        return {"paragraphs": len(paras), "runs": runs, "approx_words": words}
    except Exception:
        return None


def ext_from_filename(name: str) -> str:
    p = name.rsplit(".", 1)
    if len(p) == 2:
        return p[1].lower()
    return ""
